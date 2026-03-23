from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# Brand colors (store as hex strings for cell bg, and RGBColor for font)
NAVY_HEX = '003087'
ELECTRIC_BLUE_HEX = '00A8E0'
WHITE_HEX = 'FFFFFF'
LIGHT_BLUE_BG_HEX = 'E8F5FF'
ACCENT_TEAL_HEX = '007AB3'
DARK_TEXT_HEX = '1A1A2E'
GRAY_BG_HEX = 'F5F7FA'

NAVY = RGBColor(0x00, 0x30, 0x87)
ELECTRIC_BLUE = RGBColor(0x00, 0xA8, 0xE0)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE_BG = RGBColor(0xE8, 0xF5, 0xFF)
ACCENT_TEAL = RGBColor(0x00, 0x7A, 0xB3)
DARK_TEXT = RGBColor(0x1A, 0x1A, 0x2E)
GRAY_BG = RGBColor(0xF5, 0xF7, 0xFA)

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_run(para, text, bold=False, color=None, size=None, font_name='Noto Sans JP'):
    run = para.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    run.font.name = font_name
    # For CJK fonts
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Noto Sans JP')
    rFonts.set(qn('w:ascii'), 'Avenir Next')
    rFonts.set(qn('w:hAnsi'), 'Avenir Next')
    rPr.insert(0, rFonts)
    return run

doc = Document()

# Page setup
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)

# ===== COVER PAGE =====

# Top accent bar (simulate with table)
cover_bar = doc.add_table(rows=1, cols=1)
cover_bar.alignment = WD_TABLE_ALIGNMENT.CENTER
bar_cell = cover_bar.cell(0, 0)
set_cell_bg(bar_cell, ELECTRIC_BLUE_HEX)
bar_para = bar_cell.paragraphs[0]
bar_para.paragraph_format.space_before = Pt(4)
bar_para.paragraph_format.space_after = Pt(4)
add_run(bar_para, ' ', size=6, color=WHITE)

doc.add_paragraph()

# Cover headline
h1 = doc.add_paragraph()
h1.alignment = WD_ALIGN_PARAGRAPH.LEFT
h1.paragraph_format.space_before = Pt(20)
h1.paragraph_format.space_after = Pt(8)
add_run(h1, '薬剤抵抗性てんかんに、\nより確かな選択を。', bold=True, color=ELECTRIC_BLUE, size=28)

# Subheadline
h2 = doc.add_paragraph()
h2.alignment = WD_ALIGN_PARAGRAPH.LEFT
h2.paragraph_format.space_before = Pt(4)
h2.paragraph_format.space_after = Pt(6)
add_run(h2, '視床前核DBS（ANT-DBS）と迷走神経刺激療法（VNS）の有効性比較', bold=False, color=NAVY, size=13)

# Citation
cite = doc.add_paragraph()
cite.paragraph_format.space_after = Pt(4)
add_run(cite, 'Zhu J, et al. Journal of Clinical Neuroscience 90 (2021) 112–117', color=ACCENT_TEAL, size=9)

doc.add_paragraph()

# Key data highlight table on cover
kpi_table = doc.add_table(rows=1, cols=3)
kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
kpi_data = [
    ('65.28%', 'ANT-DBS\n発作減少率\n（術後12ヵ月）'),
    ('72.22%', 'ANT-DBS\nレスポンダー率\n（術後12ヵ月）'),
    ('22.22%', 'ANT-DBS\n発作消失率\n（術後12ヵ月）'),
]
for i, (num, label) in enumerate(kpi_data):
    cell = kpi_table.cell(0, i)
    set_cell_bg(cell, NAVY_HEX)
    cell.width = Cm(5.5)
    p_num = cell.paragraphs[0]
    p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_num.paragraph_format.space_before = Pt(10)
    add_run(p_num, num, bold=True, color=ELECTRIC_BLUE, size=22)
    p_label = cell.add_paragraph()
    p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_label.paragraph_format.space_after = Pt(10)
    add_run(p_label, label, color=WHITE, size=9)

doc.add_paragraph()

# Tagline
tl = doc.add_paragraph()
tl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(tl, 'Engineering the extraordinary', bold=False, color=NAVY, size=10, font_name='Avenir Next')

# Page break
doc.add_page_break()

# ===== SECTION 1: Disease Background =====
def add_section_header(doc, text, color=NAVY):
    # Blue left-border effect via table
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # narrow left bar
    left_cell = t.cell(0, 0)
    left_cell.width = Cm(0.35)
    set_cell_bg(left_cell, ELECTRIC_BLUE_HEX)
    p = left_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, ' ', size=6)
    # text cell
    right_cell = t.cell(0, 1)
    p2 = right_cell.paragraphs[0]
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(4)
    add_run(p2, text, bold=True, color=color, size=13)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

add_section_header(doc, 'SECTION 1　疾患背景：今もコントロールできていない患者がいる')

# Key statement box
box = doc.add_table(rows=1, cols=1)
box.alignment = WD_TABLE_ALIGNMENT.CENTER
bc = box.cell(0, 0)
set_cell_bg(bc, LIGHT_BLUE_BG_HEX)
bp = bc.paragraphs[0]
bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
bp.paragraph_format.space_before = Pt(10)
bp.paragraph_format.space_after = Pt(10)
add_run(bp, 'てんかん患者の約30%は、薬物療法では発作をコントロールできていません。', bold=True, color=NAVY, size=12)

doc.add_paragraph()

body1 = doc.add_paragraph()
body1.paragraph_format.space_after = Pt(6)
add_run(body1, '世界に約7,000万人のてんかん患者が存在し、そのうち約3割が', size=10, color=DARK_TEXT)
add_run(body1, '薬剤抵抗性てんかん（DRE）', bold=True, size=10, color=NAVY)
add_run(body1, 'に分類されます。2剤以上の抗てんかん薬（AED）で適切に治療を行っても発作が持続するこれらの患者は、切除術の適応外であるか、手術後も発作が続くケースも少なくありません。', size=10, color=DARK_TEXT)

body2 = doc.add_paragraph()
body2.paragraph_format.space_after = Pt(6)
add_run(body2, '神経刺激療法は、薬物療法・切除術に次ぐ有力な選択肢です。', bold=True, size=10, color=NAVY)

# Two-col comparison
comp_table = doc.add_table(rows=2, cols=2)
comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['VNS（迷走神経刺激療法）', 'ANT-DBS（視床前核深部脳刺激療法）']
details = ['FDA承認：1997年（4歳以上）\n薬剤抵抗性てんかんの標準的神経刺激療法', 'FDA承認：2018年（18歳以上）\nメドトロニック Activa™ システム使用']
for col, (h, d) in enumerate(zip(headers, details)):
    hc = comp_table.cell(0, col)
    set_cell_bg(hc, NAVY_HEX)
    hp = hc.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_before = Pt(6)
    hp.paragraph_format.space_after = Pt(6)
    add_run(hp, h, bold=True, color=WHITE, size=10)
    dc = comp_table.cell(1, col)
    set_cell_bg(dc, GRAY_BG_HEX)
    dp = dc.paragraphs[0]
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dp.paragraph_format.space_before = Pt(8)
    dp.paragraph_format.space_after = Pt(8)
    add_run(dp, d, size=9, color=DARK_TEXT)

doc.add_paragraph()

# ===== SECTION 2: Study Overview =====
add_section_header(doc, 'SECTION 2　試験概要：同一チームによる厳格な単施設比較')

study_intro = doc.add_paragraph()
study_intro.paragraph_format.space_after = Pt(6)
add_run(study_intro, 'これまで異なる施設・異なる評価基準で行われてきたVNSとANT-DBSの比較を、同一施設・同一専門チームにより初めて実施した後ろ向き研究です。', size=10, color=DARK_TEXT)

overview_table = doc.add_table(rows=6, cols=2)
overview_table.alignment = WD_TABLE_ALIGNMENT.CENTER
overview_data = [
    ('試験デザイン', '後ろ向き観察研究、単施設'),
    ('施設', '宣武医院 神経外科・機能神経外科センター（北京首都医科大学）'),
    ('登録期間', '2013年6月〜2018年7月'),
    ('観察期間', '術前ベースライン〜術後12ヵ月'),
    ('評価間隔', '術後3・6・9・12ヵ月'),
    ('評価方法', '外来での問診により発作頻度を記録'),
]
for r, (label, val) in enumerate(overview_data):
    lc = overview_table.cell(r, 0)
    vc = overview_table.cell(r, 1)
    bg = GRAY_BG_HEX if r % 2 == 0 else WHITE_HEX
    set_cell_bg(lc, LIGHT_BLUE_BG_HEX)
    set_cell_bg(vc, bg)
    lp = lc.paragraphs[0]
    lp.paragraph_format.space_before = Pt(5)
    lp.paragraph_format.space_after = Pt(5)
    add_run(lp, label, bold=True, color=NAVY, size=9)
    vp = vc.paragraphs[0]
    vp.paragraph_format.space_before = Pt(5)
    vp.paragraph_format.space_after = Pt(5)
    add_run(vp, val, size=9, color=DARK_TEXT)

doc.add_paragraph()

# Patient info box
pt_table = doc.add_table(rows=1, cols=2)
pt_table.alignment = WD_TABLE_ALIGNMENT.CENTER
pt_data = [
    ('VNS群  N=17', '平均年齢 20.24 ± 11.40歳\n範囲：5〜41歳\n焦点性発作：8例、全般発作：9例'),
    ('ANT-DBS群  N=18', '平均年齢 28.94 ± 12.00歳\n範囲：12〜52歳\n焦点性発作：13例、全般発作：5例'),
]
for col, (title, detail) in enumerate(pt_data):
    cell = pt_table.cell(0, col)
    set_cell_bg(cell, NAVY_HEX if col == 1 else ACCENT_TEAL_HEX)
    tp = cell.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(8)
    add_run(tp, title + '\n', bold=True, color=WHITE, size=11)
    add_run(tp, detail, color=WHITE, size=9)
    tp.paragraph_format.space_after = Pt(8)

doc.add_paragraph()

# ===== SECTION 3: Results =====
doc.add_page_break()
add_section_header(doc, 'SECTION 3　主要結果：ANT-DBSは全時点でVNSを上回る発作減少を達成')

result_headline = doc.add_paragraph()
result_headline.paragraph_format.space_after = Pt(8)
add_run(result_headline, '術後12ヵ月で、ANT-DBSはVNSを約17ポイント上回る発作減少率を示しました。', bold=True, color=NAVY, size=12)

# Efficacy table
eff_table = doc.add_table(rows=5, cols=3)
eff_table.alignment = WD_TABLE_ALIGNMENT.CENTER
eff_headers = ['経過期間', 'ANT-DBS（N=18）', 'VNS（N=17）']
eff_data = [
    ('術後3ヵ月', '57.22%', '36.06%'),
    ('術後6ヵ月', '61.61%', '39.94%'),
    ('術後9ヵ月', '63.94%', '45.24%'),
    ('術後12ヵ月 ★', '65.28%', '48.35%'),
]
# Header row
for col, h in enumerate(eff_headers):
    cell = eff_table.cell(0, col)
    set_cell_bg(cell, NAVY_HEX)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    add_run(p, h, bold=True, color=WHITE, size=10)

for r, (period, dbs, vns) in enumerate(eff_data):
    is_last = r == 3
    row_bg = LIGHT_BLUE_BG_HEX if is_last else (GRAY_BG_HEX if r % 2 == 0 else WHITE_HEX)
    for col, val in enumerate([period, dbs, vns]):
        cell = eff_table.cell(r + 1, col)
        set_cell_bg(cell, row_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        bold = is_last or col == 1
        color = ELECTRIC_BLUE if (col == 1 and is_last) else NAVY if col == 1 else DARK_TEXT
        add_run(p, val, bold=bold, color=color, size=11 if (col == 1 and is_last) else 10)

doc.add_paragraph()

# Responder rate table
resp_label = doc.add_paragraph()
resp_label.paragraph_format.space_after = Pt(4)
add_run(resp_label, 'レスポンダー率・発作消失率（術後12ヵ月）', bold=True, color=NAVY, size=11)

resp_table = doc.add_table(rows=4, cols=3)
resp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
resp_headers = ['指標', 'ANT-DBS（N=18）', 'VNS（N=17）']
resp_data = [
    ('レスポンダー率\n（発作50%以上減少）', '72.22%  (N=13)', '58.82%  (N=10)'),
    ('発作消失率', '22.22%  (N=4)', '17.65%  (N=3)'),
    ('非レスポンダー', '27.78%  (N=5)', '41.18%  (N=7)'),
]
for col, h in enumerate(resp_headers):
    cell = resp_table.cell(0, col)
    set_cell_bg(cell, NAVY_HEX)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    add_run(p, h, bold=True, color=WHITE, size=10)

for r, (metric, dbs_val, vns_val) in enumerate(resp_data):
    bg = GRAY_BG_HEX if r % 2 == 0 else WHITE_HEX
    for col, val in enumerate([metric, dbs_val, vns_val]):
        cell = resp_table.cell(r + 1, col)
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        color = ELECTRIC_BLUE if col == 1 else DARK_TEXT
        add_run(p, val, bold=(col == 1), color=color, size=10)

doc.add_paragraph()

# ===== SECTION 4: Patient Selection =====
add_section_header(doc, 'SECTION 4　患者選択の手引き：術前データから治療選択の根拠を')

sel_intro = doc.add_paragraph()
sel_intro.paragraph_format.space_after = Pt(8)
add_run(sel_intro, '本試験の結果は、術前の臨床的特徴に基づいた治療選択の可能性を示しています。', size=10, color=DARK_TEXT)

sel_table = doc.add_table(rows=2, cols=2)
sel_table.alignment = WD_TABLE_ALIGNMENT.CENTER

sel_data = [
    ('ANT-DBSが適している可能性が高い患者', [
        '・手術時年齢が高い',
        '・焦点性発作（発作減少率 71.15%）',
        '・開頭術の既往がある（発作減少率 81.20%）',
        '・てんかん罹病期間が長い（P<0.05）',
    ]),
    ('VNSが適している可能性が高い患者', [
        '・発症年齢が高い（P<0.05）',
        '・手術時年齢が高い',
        '・焦点性発作（発作減少率 59.00%）',
    ]),
]

for col, (title, items) in enumerate(sel_data):
    hcell = sel_table.cell(0, col)
    set_cell_bg(hcell, NAVY_HEX if col == 0 else ACCENT_TEAL_HEX)
    hp = hcell.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_before = Pt(8)
    hp.paragraph_format.space_after = Pt(8)
    add_run(hp, title, bold=True, color=WHITE, size=10)

    dcell = sel_table.cell(1, col)
    set_cell_bg(dcell, LIGHT_BLUE_BG_HEX if col == 0 else GRAY_BG_HEX)
    dp = dcell.paragraphs[0]
    dp.paragraph_format.space_before = Pt(8)
    dp.paragraph_format.space_after = Pt(8)
    add_run(dp, '\n'.join(items), size=9, color=DARK_TEXT)

doc.add_paragraph()

# ===== SECTION 5: Safety =====
add_section_header(doc, 'SECTION 5　安全性：継続刺激による良好な忍容性')

safety_box = doc.add_table(rows=1, cols=1)
safety_box.alignment = WD_TABLE_ALIGNMENT.CENTER
sc = safety_box.cell(0, 0)
set_cell_bg(sc, LIGHT_BLUE_BG_HEX)
sp = sc.paragraphs[0]
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp.paragraph_format.space_before = Pt(10)
sp.paragraph_format.space_after = Pt(10)
add_run(sp, '観察期間中、両治療とも重篤な合併症・副作用は報告されませんでした。', bold=True, color=NAVY, size=11)

doc.add_paragraph()

safety_table = doc.add_table(rows=3, cols=2)
safety_table.alignment = WD_TABLE_ALIGNMENT.CENTER
safe_headers = ['治療', '主な有害事象']
safe_data = [
    ('VNS', '嗄声、咳嗽、疼痛、呼吸困難'),
    ('ANT-DBS', '植込み部疼痛、感染、リード偏位'),
]
for col, h in enumerate(safe_headers):
    cell = safety_table.cell(0, col)
    set_cell_bg(cell, NAVY_HEX)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    add_run(p, h, bold=True, color=WHITE, size=10)

for r, (tx, ae) in enumerate(safe_data):
    bg = GRAY_BG_HEX if r % 2 == 0 else WHITE_HEX
    for col, val in enumerate([tx, ae]):
        cell = safety_table.cell(r + 1, col)
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        add_run(p, val, size=10, color=DARK_TEXT)

doc.add_paragraph()

# ===== CONCLUSION =====
add_section_header(doc, 'まとめ')

conc_table = doc.add_table(rows=1, cols=3)
conc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
conc_items = [
    ('01', 'ANT-DBSは全時点でVNSを上回る発作減少率を達成'),
    ('02', '術前の臨床データで治療選択の根拠を得られる可能性'),
    ('03', '両治療とも時間経過で効果が蓄積し、安全に継続可能'),
]
for col, (num, text) in enumerate(conc_items):
    cell = conc_table.cell(0, col)
    set_cell_bg(cell, NAVY_HEX)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    add_run(p, num + '\n', bold=True, color=ELECTRIC_BLUE, size=18)
    add_run(p, text, color=WHITE, size=9)
    p.paragraph_format.space_after = Pt(10)

doc.add_paragraph()

# ===== CALL TO ACTION =====
cta_table = doc.add_table(rows=1, cols=1)
cta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cc = cta_table.cell(0, 0)
set_cell_bg(cc, ELECTRIC_BLUE_HEX)
cp = cc.paragraphs[0]
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
cp.paragraph_format.space_before = Pt(12)
add_run(cp, '薬剤抵抗性てんかんの患者さんに、次のステップを。\n', bold=True, color=WHITE, size=14)
add_run(cp, 'ANT-DBSの治療選択・デバイス情報は、メドトロニック担当者へお問い合わせください。', color=WHITE, size=10)
cp.paragraph_format.space_after = Pt(12)

doc.add_paragraph()

# Footer / disclaimer
footer_p = doc.add_paragraph()
footer_p.paragraph_format.space_before = Pt(8)
add_run(footer_p, 'Medtronic Activa™ PC（型番：37601）/ Activa™ RC（型番：37612）\n', bold=True, color=NAVY, size=9)
add_run(footer_p, '本資材は医療従事者向けの情報提供を目的として作成されています。記載の臨床データはZhu J, et al. J Clin Neurosci. 2021;90:112–117 を出典とします。\n', size=8, color=RGBColor(0x66, 0x66, 0x66))

tagline_p = doc.add_paragraph()
tagline_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(tagline_p, 'Engineering the extraordinary   |   Medtronic', bold=False, color=NAVY, size=9, font_name='Avenir Next')

# Save
output_path = '/home/user/Shota/HCP_ANT-DBS_EvidenceSummary.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
